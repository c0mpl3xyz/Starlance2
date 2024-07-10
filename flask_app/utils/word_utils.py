from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
import os
import uuid

TEMPLATE_PATH = os.getenv('TEMPLATE_PATH')

def replace_bookmarks(doc_path, replacements, tables_rows=None):
    try:
        doc = Document(doc_path)

        for paragraph in doc.paragraphs:
            for bookmark_name in replacements:
                if f'<{bookmark_name}>' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'<{bookmark_name}>', str(replacements[bookmark_name]))
        
        if tables_rows is not None and len(tables_rows):
            fill_tables(doc, tables_rows)

        new_file_path = os.path.join(TEMPLATE_PATH, f'{str(uuid.uuid1())}.docx')
        doc.save(new_file_path)
        return True, new_file_path

    except Exception as e:
        print(f"Error replacing bookmarks: {e}")
        raise e
        return False, None

def fill_tables(doc, table_data) -> bool:
    if not table_data:
        return False

    # Find the paragraph containing the <TABLE> placeholder
    table_found = False
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.left_indent = Mm(30.4)
        paragraph.paragraph_format.right_indent = Mm(14.4)
        if '<TABLE>' in paragraph.text:
            table_found = True
            # Remove the old paragraph containing the bookmark
            p = paragraph._element
            p.getparent().remove(p)

            # Create a new table
            table = doc.add_table(rows=1, cols=len(table_data[0]))  # Assuming table_data[0] contains headers
            table.autofit = False  # Disable autofit to set fixed widths

            # Define the width for each column
            section = doc.sections[0] if doc.sections else None
            block_width = None
            if section and section.page_width is not None and section.left_margin is not None and section.right_margin is not None:
                block_width = section.page_width - section.left_margin - section.right_margin
            else:
                block_width = Inches(6)  # Default width in Inches (adjust as needed)

            column_widths = [block_width / len(table_data[0]) for _ in range(len(table_data[0]))]

            # Add table headers dynamically from the first row keys
            hdr_cells = table.rows[0].cells
            headers = list(table_data[0].keys())
            for i, header_text in enumerate(headers):
                hdr_cells[i].width = column_widths[i]
                hdr_cells[i].text = header_text
                hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = hdr_cells[i].paragraphs[0].runs[0]
                font = run.font
                font.size = Pt(10)
                font.bold = True

            # Add table rows
            for row_data in table_data:
                row_cells = table.add_row().cells
                for i, cell_value in enumerate(row_data.values()):
                    row_cells[i].width = column_widths[i]
                    row_cells[i].text = str(cell_value)
                    row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = row_cells[i].paragraphs[0].runs[0]
                    font = run.font
                    font.size = Pt(10)

    if not table_found:
        print("Error: <TABLE> placeholder not found in the document.")
    return True